import logging
import os
from typing import BinaryIO, List, Dict, Any, Optional
from io import BytesIO

import fitz  # PyMuPDF
from langchain.schema import Document

logger = logging.getLogger(__name__)

class DocumentParser:
    """Parser for PDF and other document types"""
    
    def parse_pdf(self, file: BinaryIO, metadata: Dict[str, Any] = None) -> List[Document]:
        """
        Parse a PDF file into Document objects
        
        Args:
            file: File-like object containing PDF data
            metadata: Optional metadata to add to the documents
            
        Returns:
            List of Document objects
        """
        try:
            # Copy to BytesIO to ensure seekable
            buffer = BytesIO(file.read())
            
            # Open PDF with PyMuPDF
            logger.info("Parsing PDF document")
            pdf = fitz.open(stream=buffer, filetype="pdf")
            
            # Get PDF metadata
            pdf_metadata = {
                "title": pdf.metadata.get("title", ""),
                "author": pdf.metadata.get("author", ""),
                "subject": pdf.metadata.get("subject", ""),
                "keywords": pdf.metadata.get("keywords", ""),
                "creator": pdf.metadata.get("creator", ""),
                "producer": pdf.metadata.get("producer", ""),
                "page_count": len(pdf),
                "format": "pdf"
            }
            
            # Combine with passed-in metadata
            if metadata:
                pdf_metadata.update(metadata)
            
            # Extract text from each page
            documents = []
            for page_num, page in enumerate(pdf):
                text = page.get_text()
                
                # Remove NULL characters that can cause PostgreSQL errors
                text = text.replace('\x00', '')
                
                # Skip empty pages
                if not text.strip():
                    continue
                
                # Create a document for this page
                page_metadata = pdf_metadata.copy()
                page_metadata["page_num"] = page_num + 1
                
                doc = Document(
                    page_content=text,
                    metadata=page_metadata
                )
                documents.append(doc)
            
            logger.info(f"Parsed {len(documents)} pages from PDF")
            return documents
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            return []
        
    def parse_docx(self, file: BinaryIO, metadata: Dict[str, Any] = None) -> List[Document]:
        """
        Parse a DOCX file into Document objects
        
        Args:
            file: File-like object containing DOCX data
            metadata: Optional metadata to add to the documents
            
        Returns:
            List of Document objects
        """
        try:
            # For DOCX, we need python-docx
            import docx
            
            # Copy to BytesIO to ensure seekable
            buffer = BytesIO(file.read())
            
            # Open DOCX
            logger.info("Parsing DOCX document")
            doc = docx.Document(buffer)
            
            # Extract text from paragraphs
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            
            # Create document metadata
            doc_metadata = {
                "title": doc.core_properties.title if hasattr(doc, 'core_properties') else "",
                "author": doc.core_properties.author if hasattr(doc, 'core_properties') else "",
                "format": "docx"
            }
            
            # Combine with passed-in metadata
            if metadata:
                doc_metadata.update(metadata)
            
            # Create a document
            document = Document(
                page_content=full_text,
                metadata=doc_metadata
            )
            
            logger.info(f"Parsed DOCX document")
            return [document]
            
        except ImportError:
            logger.error("python-docx not installed, cannot parse DOCX files")
            return []
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            return []
    
    def parse(self, file: BinaryIO, metadata: Dict[str, Any] = None) -> List[Document]:
        """
        Parse a document into Document objects based on the file type
        
        Args:
            file: File-like object
            metadata: Optional metadata to add to the documents
            
        Returns:
            List of Document objects
        """
        # Determine file type from metadata
        file_format = metadata.get("mime_type", "") if metadata else ""
        
        if file_format == "application/pdf" or file_format.endswith("/pdf"):
            return self.parse_pdf(file, metadata)
        elif file_format == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self.parse_docx(file, metadata)
        else:
            logger.warning(f"Unsupported document format: {file_format}")
            return []

# Singleton instance
document_parser = DocumentParser()
