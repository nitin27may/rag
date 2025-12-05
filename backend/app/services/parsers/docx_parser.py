import logging
from typing import List, Dict, Any, BinaryIO, Optional
import io

from langchain.schema import Document
import docx

from app.services.parsers.base import BaseParser
from app.services.chunking_service import chunking_service

logger = logging.getLogger(__name__)

class DocxParser(BaseParser):
    """Parser for Microsoft Word documents"""
    
    def __init__(self):
        super().__init__()
        self.supported_mimetypes = [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]
        # Use centralized chunking service
        self.chunking_service = chunking_service
    
    def parse(self, file_data: BinaryIO, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """
        Parse a DOCX document into chunks
        
        Args:
            file_data: File-like object containing the DOCX data
            metadata: Optional metadata to include with the document
            
        Returns:
            List of Document objects
        """
        if metadata is None:
            metadata = {}
        
        try:
            # Read the document
            doc = docx.Document(file_data)
            
            # Add document metadata
            metadata["format"] = "docx"
            metadata["paragraph_count"] = len(doc.paragraphs)
            
            if doc.core_properties:
                props = doc.core_properties
                if props.title:
                    metadata["title"] = props.title
                if props.author:
                    metadata["author"] = props.author
                if props.subject:
                    metadata["subject"] = props.subject
                if props.keywords:
                    metadata["keywords"] = props.keywords
            
            # Extract text from the document
            full_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            if full_text:
                doc = Document(
                    page_content=full_text,
                    metadata=metadata
                )
                
                # Split into chunks using centralized chunking service
                return self.chunking_service.split_documents([doc])
            
            return []
        
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise